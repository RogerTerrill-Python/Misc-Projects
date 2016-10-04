#!/usr/bin/env python3

import os, docx

from docx import Document

findText = input("Type text to replace: ") 

for dirpath, dirnames, filenames in os.walk('.'):
    printname = [fn for fn in filenames if fn.endswith('.docx')]
    print(printname)
    for name in printname:
        print('List Name: ' + name)
    
    
    for basename in (fn for fn in filenames if fn.endswith('.docx')):
        filename = os.path.join(dirpath, basename)
        document = docx.Document(filename)
        print('Current Directory is: %s' % dirpath)
        print('basename is: %s' % basename)
        print('filename is:' + filename + '\n')
        
        # Check through all the tables for text
        tables = document.tables
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if findText in paragraph.text:                              
                            inline = paragraph.runs                                 
                            for i in range(len(inline)):
                                if findText in inline[i].text:
                                    text = inline[i].text.replace(findText, '')
                                    inline[i].text = text
 
         # Check through all paragraphs for word                           
        for paragraph in document.paragraphs:                           
            if findText in paragraph.text:                              
                inline = paragraph.runs                                 
                for i in range(len(inline)):
                    if findText in inline[i].text:
                        text = inline[i].text.replace(findText, '')
                        inline[i].text = text
        # Save file 
        document.save(filename)  